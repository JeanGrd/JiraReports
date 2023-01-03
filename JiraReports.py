# -*- coding: utf-8 -*-
"""
---------------------------------------------------------------------------------------------------------------
Title: Transform an XML document to a pandas dataframe to generate Jira reports more quickly

Remark : The main objective is to generate reports containing tables that contain jira data in a Word or Excel
format. These tables are easily defined in XML documents (for more information : Readme.md).

Written by: Jean Guiraud
---------------------------------------------------------------------------------------------------------------
"""

# !/usr/bin/env python3
# Python version : 3.8

import os
import sys
import xml.etree.ElementTree as ET
import pandas as pd
from jira.client import JIRA
from tqdm import tqdm

import docx
from docx.oxml.shared import qn  # Feel free to move these out
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


class JiraReports:
    """
    Transform XML Document from Jira server to a pandas dataframe

    :param jira: jira instance
    :param xml: link to the XML document
    """

    def __init__(self, jira: JIRA, xml: str):

        file = ET.parse(xml)  # To parse the content in a variable
        self.__root = file.getroot()  # To go to the beginning of the XML document
        self.file = []

        for table in tqdm(self.__root.findall("Table")):

            tab = {}
            for incremental, columns in enumerate(table.findall("Column")):
                tab[incremental] = [columns.get("type"), columns.text]

            if table.get("style") == "Classic":
                jira_issues = jira.search_issues(table.find("JQL").text, maxResults=False)
                self.file.append(pd.DataFrame(jira_import(jira_issues, tab)))

            elif table.get("style") == "MultipleJQL":

                pandas_tables = []

                for filters in table.findall("Filters"):
                    for nbJQL in filters.findall("JQL"):
                        jira_issues = jira.search_issues(nbJQL.text, maxResults=False)
                        pandas_tables.append(pd.DataFrame(jira_import(jira_issues, tab)))

                self.file.append(pandas_tables)

            elif table.get("style") == "LinkOneTicket":
                jira_issues = jira.search_issues(table.find("JQL").text, maxResults=1)

                jira_issues = jira.search_issues(
                    'issue in linkedIssues(' + jira_issues[0].key + ', ' +
                    table.find("JQL").get("link") + ')', maxResults=False)

                self.file.append(pd.DataFrame(jira_import(jira_issues, tab)))

    def to_excel(self, document_name: str = "jira_excel"):
        """
        Generate Excel document

        :param document_name: the document name *(default : jira_excel)*
                / you can add the path, like : users/john/test
        """

        writer = pd.ExcelWriter(document_name + ".xlsx", engine="xlsxwriter")
        workbook = writer.book  # Creating an excel document

        # Formatting of the excel document
        header = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#D8E4BC'})
        Bold = workbook.add_format({'bold': True, 'align': 'center'})
        # To give the title of the document according to the XML file
        titre = workbook.add_format({'align': 'center', 'bold': True})

        worksheet_name = []

        for incremental, all_tables in enumerate(self.__root.findall('Table')):

            Name = all_tables.get("name")
            tag = ["/", "*", ":", "[", "]"]

            for tags in tag:  # For prohibited characters
                Name = Name.replace(tags, ' ')

            Name = str(incremental + 1) + " - " + Name

            if len(Name) > 31:  # For character length
                Name = Name[:31]

            worksheet_name.append(Name)

            if all_tables.get('style') == "Classic" or all_tables.get('style') == "LinkOneTicket":
                self.file[incremental].to_excel(writer, sheet_name=Name, startrow=3, header=False, index=False)

            elif all_tables.get('style') == "MultipleJQL":
                start_r = 3
                for inc_JQL, nbJQL in enumerate(all_tables[0].findall("JQL")):
                    self.file[incremental][inc_JQL].to_excel(writer, sheet_name=Name, startrow=start_r + 1,
                                                             header=False, index=False)

                    writer.sheets[Name].merge_range(start_r, 0, start_r, len(all_tables.findall('Column')) - 1,
                                                    'Merged Range')
                    writer.sheets[Name].write(start_r, 0, nbJQL.get("name"), Bold)

                    start_r += len(self.file[incremental][inc_JQL].index)

            # To put the name on the Excel sheet
            writer.sheets[Name].write(0, 0, all_tables.get('name'), titre)

            for incremental, column in enumerate(all_tables.findall('Column')):  # Creation of all the columns
                writer.sheets[Name].write(2, incremental, column.get('name'), header)

            writer.sheets[Name].set_row(2, 30)

        writer.close()

    def to_word(self, document_name: str = "jira_word", landscape: bool = False, cell_color: str = "#85B1ED"):
        """
        Generate Word document from Jira_XMLDocument instance

        :param document_name: the document name (default : jira_word)
        :param landscape: document in landscape format (True or False)
        :param cell_color:
        """

        document = docx.Document()
        document.add_heading(self.__root.get('name'), level=0)

        if landscape:
            new_width, new_height = document.sections[-1].page_height, document.sections[-1].page_width
            document.sections[-1].page_width = new_width
            document.sections[-1].page_height = new_height

        for incremental, all_tables in enumerate(self.__root.findall('Table')):

            document.add_heading(all_tables.get('name'), level=1)
            table = document.add_table(rows=0, cols=len(all_tables.findall('Column')))
            p = table.add_row().cells

            for num_cell, column in enumerate(all_tables.findall('Column')):
                table.cell(0, num_cell).text = column.get('name')
                table.cell(0, num_cell).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table.rows[0].cells[num_cell].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                __set_cell_background__(table.rows[0].cells[num_cell], cell_color)

            __make_rows_bold__(table.rows[0])
            __set_repeat_table_header__(table.rows[0])
            table.style = 'Table Grid'
            table.add_row()

            if all_tables.get("style") == "Classic" or all_tables.get('style') == "LinkOneTicket":

                # add the rest of the data frame
                for i in range(self.file[incremental].shape[0]):
                    if i != 0:
                        table.add_row()
                    for j in range(self.file[incremental].shape[-1]):
                        table.cell(i + 1, j).text = str(self.file[incremental].values[i, j])
                        table.cell(i + 1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            elif all_tables.get("style") == "MultipleJQL":
                row = 1

                for inc_JQL, nbJQL in enumerate(all_tables[0].findall("JQL")):

                    if inc_JQL != 0:
                        table.add_row()

                    table.cell(row, 0).text = nbJQL.get("name")
                    __make_rows_bold__(table.rows[row])

                    for column in range(self.file[incremental][inc_JQL].shape[-1]):
                        table.cell(row, 0).merge(table.cell(row, column))

                    table.cell(row, 0).paragraphs[
                        0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    row += 1

                    # add the rest of the data frame
                    for i in range(self.file[incremental][inc_JQL].shape[0]):
                        table.add_row()
                        for j in range(self.file[incremental][inc_JQL].shape[-1]):
                            table.cell(row, j).text = str(
                                self.file[incremental][inc_JQL].values[i, j])
                            table.cell(i + 1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        row += 1
            # save the doc
        document.save(document_name + '.docx')

    def to_word_template(self, path_template_word: str, document_name: str = "jira_word_template"):
        """
        Generate Word template document from Jira_XMLDocument instance

        :param path_template_word: path of the word template *(see Readme.md)*
        :param document_name: the document name *(default : jira_word_template)*
                / you can add the path, like : users/john/test
        """

        # open an existing document
        document = docx.Document(path_template_word)
        tables = document.tables

        for incremental, all_tables in enumerate(self.__root.findall('Table')):

            for search in range(len(tables)):
                try:
                    if all_tables.get("keyword") == tables[search].cell(1, 0).text:
                        tables_length = search
                except:
                    continue

                search += 1

            if 'tables_length' in locals():

                if all_tables.get("style") == "Classic" or all_tables.get('style') == "LinkOneTicket":
                    # add the rest of the data frame
                    for i in range(self.file[incremental].shape[0]):
                        if i != 0:
                            tables[0].add_row()
                        for j in range(self.file[incremental].shape[-1]):
                            tables[tables_length].cell(i + 1, j).text = str(self.file[incremental].values[i, j])

                elif all_tables.get("style") == "MultipleJQL":

                    row = 1

                    for inc_JQL, nbJQL in enumerate(all_tables[0].findall("JQL")):

                        if row != 1:
                            tables[tables_length].add_row()

                        tables[tables_length].cell(row, 0).text = nbJQL.get("name")
                        __make_rows_bold__(tables[tables_length].rows[row])

                        for column in range(self.file[incremental][inc_JQL].shape[-1]):
                            tables[tables_length].cell(row, 0).merge(tables[tables_length].cell(row, column))

                        tables[tables_length].cell(row, 0).paragraphs[
                            0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                        row += 1

                        # add the rest of the data frame
                        for i in range(self.file[incremental][inc_JQL].shape[0]):
                            tables[tables_length].add_row()
                            for j in range(self.file[incremental][inc_JQL].shape[-1]):
                                tables[tables_length].cell(row, j).text = str(
                                    self.file[incremental][inc_JQL].values[i, j])
                            row += 1
        # save the doc
        document.save(document_name + '.docx')

    def __str__(self) -> str:
        return "Jira_XMLDocument : " + self.__root.attrib["name"]


def jira_import(jira_issues: dict, information: dict) -> list:
    """
    Extract data from jira to a classic list

    :param jira_issues; jira ResultList[] *(dictionary)*
    :param information: dictionary that contains 1 key and 2 values:
            - Key = index of the column (1, 2, 3, etc...)
            - Value 1 = data to extract (example : summary, customfield, description, etc...)
            - Value 2 = type of data (example : multiplevalue, link, etc...)
            example : {0: ['', 'summary'], 1: ['', 'summary'], 2: ['', 'description']}
    """

    import_array = []

    if len(jira_issues) != 0:

        for incremental, all_issues in enumerate(jira_issues):  # Browse all issues in the JSON file
            import_array.append({})

            for inc, table in information.items():

                # Condition for the multiple values in the JSON variable
                if table[0] == "multiple_values":  # To check that it is called in the XML file

                    Fulltext = ""  # Variable that adds the names of the multiple values
                    for multiple_values in eval("all_issues.fields." + table[1]):
                        if Fulltext is None:
                            Fulltext = str(multiple_values.name)
                            continue
                        Fulltext += "\n" + str(multiple_values.name)

                    import_array[incremental][table[0]] = Fulltext

                elif table[0] == "link":

                    Fulltext = []

                    for link in all_issues.fields.issuelinks:
                        if link.type.inward == table[1]:
                            if hasattr(link, "inwardIssue"):
                                cut_string = str(link.inwardIssue.fields.summary).split(
                                    "_")  # To get the name of the issue

                                Fulltext.append(str(cut_string[0]) + " Iss. " + str(cut_string[1]))

                        elif link.type.outward == table[1]:
                            if hasattr(link, "outwardIssue"):
                                cut_string = str(link.outwardIssue.fields.summary).split(
                                    "_")  # To get the name of the issue

                                Fulltext.append(str(cut_string[0]) + " Iss. " + str(cut_string[1]))

                    Text = ""
                    for Fulltext in Fulltext:
                        if Text is None:
                            Text = Fulltext
                            continue
                        Text += "\n" + Fulltext

                    import_array[incremental][table[1]] = Text

                elif table[0] == "specific_summary":
                    cut_string = str(all_issues.fields.summary).split("_")  # Split variable
                    import_array[incremental][table[1]] = cut_string[2]

                # For all other fields
                else:
                    value = table[1]
                    if value in import_array[incremental]:
                        value += str(inc)
                    import_array[incremental][value] = str(eval("all_issues.fields." + table[1]))

    return import_array


def __make_rows_bold__(*rows):
    """
    Set a row in bold

    :param *rows: python-docx row
    """

    for row in rows:  # Select all rows
        for cell in row.cells:  # Select all cells
            for paragraph in cell.paragraphs:  # Select all cell's paragraph
                for run in paragraph.runs:  # Select all paragraph's run
                    run.font.bold = True  # Set in bold the run


def __set_repeat_table_header__(row):
    """
    Set repeat table row on every new page

    :param row: python-docx row
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


def __set_cell_background__(cell, fill):
    """
    Set color background in a cellular

    :param cell: xlsxwriter cell
    :param fill: specifies the color to be used for the background
    """

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # In case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd')  # Add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # Set fill property, respecting namespace
    cell_properties.append(cell_shading)  # Finally extend cell props with shading element
