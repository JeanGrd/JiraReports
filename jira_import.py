# -*- coding: utf-8 -*-
"""

---------------------------------------------------------------------------------------------------------------

Remark : This code allows to import jira data which are in JSON format to a pandas
dataframe. This part of program has been developed for the Moko application.

Written by : Jean Guiraud

---------------------------------------------------------------------------------------------------------------

"""

# !/usr/bin/env python3
# Python version : 3.8

import os
import sys
import xml.etree.ElementTree as ET

import docx
import pandas as pd
from jira.client import JIRA
from tqdm import tqdm

#import win32com.client as win32

class Jira_XMLDocument():

    """
    Transform XML Document from Jira server to a pandas dataframe

    **params**
        - **jira** = jira instance
        - **xml** = link to the XML document

    """

    def __init__(self, jira: JIRA, xml: str):

        file = ET.parse(xml)  # To parse the content in a variable
        self.root = file.getroot()  # To go to the beginning of the XML document

        # Initialization of the current path
        if getattr(sys, 'frozen', False):
            self.application_path = os.path.dirname(sys.executable)  # If the application was launched from an exe file
        elif __file__:
            self.application_path = os.path.dirname(__file__)  # If the application was launched from a py file

        self.file = []
        for table in tqdm(self.root.findall("Table")):

            tab = {}
            for incremental, columns in enumerate(table.findall("Column")):
                tab[incremental] = [columns.get("type"), columns.text]

            if table.get("style") == "Classic":
                jira_issues = jira.search_issues(table.find("JQL").text, maxResults=False)
                self.file.append(pd.DataFrame(self.__jira_import__(jira_issues, tab)))

            elif table.get("style") == "MultipleFilters":

                pandas_tables = []

                for filters in table.findall("Filters"):
                    for nbJQL in filters.findall("JQL"):
                        jira_issues = jira.search_issues(nbJQL.text, maxResults=False)
                        pandas_tables.append(pd.DataFrame(self.__jira_import__(jira_issues, tab)))

                self.file.append(pandas_tables)

            elif table.get("style") == "LinkOneTicket":
                jira_issues = jira.search_issues(table.find("JQL").text, maxResults=1)

                jira_issues = jira.search_issues(
                    'issue in linkedIssues(' + jira_issues[0].key + ', ' +
                    table.find("JQL").get("link") + ')', maxResults=False)

                self.file.append(pd.DataFrame(self.__jira_import__(jira_issues, tab)))

    def to_excel(self, docname="jira_excel", path=""):

        """

        Generate Excel document

        **params**
            - **docname** = the document name *(default : jira_excel)*
            - **path** = the path where the file will be saved *(default : "")*

        """

        writer = pd.ExcelWriter(path + docname + ".xlsx", engine="xlsxwriter")
        workbook = writer.book  # Creating an excel document

        # Formatting of the excel document
        header = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#D8E4BC'})
        Bold = workbook.add_format({'bold': True, 'align': 'center'})
        # To give the title of the document according to the XML file
        titre = workbook.add_format({'align': 'center', 'bold': True})

        worksheet_name = []

        for incremental, all_tables in enumerate(self.root.findall('Table')):

            Name = all_tables.get("name")
            tag = ["/", "*", ":", "[", "]"]

            for tags in tag:  # For prohibited characters
                Name = Name.replace(tags, ' ')

            Name = str(incremental + 1) + " - " + Name

            if len(Name) > 31:  # For character length
                Name = Name[:31]

            worksheet_name.append(Name)

            if all_tables.get('style') == "Classic" or all_tables.get('style') == "LinkOneTicket":
                self.file[incremental].to_excel(writer, sheet_name=Name, startrow=3, header=False, index=False)

            elif all_tables.get('style') == "MultipleFilters":
                start_r = 4
                for filters in all_tables.findall("Filters"):
                    for inc_JQL, nbJQL in enumerate(filters.findall("JQL")):
                        self.file[incremental][inc_JQL].to_excel(writer, sheet_name=Name, startrow=start_r,
                                                                 header=False, index=False)

                        writer.sheets[Name].merge_range(start_r - 1, 0, start_r - 1, len(all_tables.findall('Column')),
                                                        'Merged Range')
                        writer.sheets[Name].write(start_r - 1, 0, nbJQL.get("name"), Bold)

                        start_r += len(self.file[incremental][inc_JQL].index) + 1

            # To put the name on the Excel sheet
            writer.sheets[Name].write(0, 0, all_tables.get('name'), titre)

            for incremental, column in enumerate(all_tables.findall('Column')):  # Creation of all the columns
                writer.sheets[Name].write(2, incremental, column.get('name'), header)

            writer.sheets[Name].set_row(2, 30)

        writer.close()

        """
        excel = win32.gencache.EnsureDispatch('Excel.Application')
        wb = excel.Workbooks.Open(self.application_path + "/" + path + docname + ".xlsx")

        for name in worksheet_name:
            ws = wb.Worksheets(name)
            ws.Columns.AutoFit()

        wb.Save()
        wb.Close(True)
        """

    def to_word(self):

        """

        Generate Word document

        **params**
            - **docname** = the document name *(default : jira_word)*
            - **path** = the path where the file will be saved *(default : None)*

        """

        pass  # TODO

    def to_word_template(self, path_template_word: str, docname="jira_word_template", path=""):

        """

        Generate Word template document

        **params**
            - **path_template_word = path of the word template *(see Readme.md)*
            - **docname** = the document name *(default : jira_word_template)*
            - **path** = the path where the file will be saved *(default : "")*

        """

        # open an existing document
        document = docx.Document(path_template_word)
        tables = document.tables

        for incremental, all_tables in enumerate(self.root.findall('Table')):

            for searchtables in range(len(tables)):
                try:
                    if all_tables.get("keyword") == tables[searchtables].cell(1, 0).text:
                        tableslenght = searchtables
                except:
                    continue

                searchtables += 1

            if all_tables.get("style") == "Classic" or all_tables.get('style') == "LinkOneTicket":

                # add the rest of the data frame
                for i in range(self.file[incremental].shape[0]):
                    if i != 0:
                        tables[tableslenght].add_row()
                    for j in range(self.file[incremental].shape[-1]):
                        tables[tableslenght].cell(i + 1, j).text = str(self.file[incremental].values[i, j])

            elif all_tables.get("style") == "MultipleFilters":

                row = 1

                for filters in all_tables.findall("Filters"):  # TODO simplify
                    for inc_JQL, nbJQL in enumerate(filters.findall("JQL")):

                        if row != 1:
                            tables[tableslenght].add_row()

                        tables[tableslenght].cell(row, 0).text = nbJQL.get("name")
                        self.__make_rows_bold__(tables[tableslenght].rows[row])

                        for column in range(self.file[incremental][inc_JQL].shape[-1]):
                            tables[tableslenght].cell(row, 0).merge(tables[tableslenght].cell(row, column))

                        tables[tableslenght].cell(row, 0).paragraphs[
                            0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                        row += 1

                        # add the rest of the data frame
                        for i in range(self.file[incremental][inc_JQL].shape[0]):
                            tables[tableslenght].add_row()
                            for j in range(self.file[incremental][inc_JQL].shape[-1]):
                                tables[tableslenght].cell(row, j).text = str(
                                    self.file[incremental][inc_JQL].values[i, j])
                            row += 1
        # save the doc
        document.save(path + docname + '.docx')

    def __make_rows_bold__(self, *rows):

        """

        Set a row in bold

        **param**
            - **row** = python-docx row

        """

        for row in rows:  # Select all rows
            for cell in row.cells:  # Select all cells
                for paragraph in cell.paragraphs:  # Select all cell's paragraph
                    for run in paragraph.runs:  # Select all paragraph's run
                        run.font.bold = True  # Set in bold the run

    def __just_highest_issues__ (self, splitter: str, n_splitter: str, version: int, list_to_split: list) -> list:

        """

        Keep the highest issue from a list where issue is located in a text box with separators

        **params**
            - **splitter** = the separator
            - **n_splitter** = the position of the issue name *(example : issue_2 the position is 0 : issue)*
            - **version** = the position of the version that you want to compare *(example issue_2 is the position 1 : 2)*
            - **list_to_split** = your list where you want to keep the highest issue

        """

        list_to_split.sort()
        list_highest_issue = []

        val = 0
        while val < len(list_to_split) - 1:
            max = 0
            groupby = element = list_to_split[val].split(splitter)

            while element[n_splitter] == groupby[n_splitter]:

                if max < int(element[version]):
                    max = int(element[version])
                    maxInd = val

                val += 1

                if val > len(list_to_split) - 1:
                    break

                element = list_to_split[val].split(splitter)

            list_highest_issue.append(list_to_split[maxInd])

        return list_highest_issue

    def __jira_import__(self, jira_issues: dict, information: dict) -> list:

        """

        Extract data from jira to a classic list

        **params**
            1. **jira_issues** = jira ResultList[] *(dictionary)*
            2. **information** = dictionary that contains 1 key and 2 values:
                - *Key*, index of the column (1, 2, 3, etc...)
                - *Value 1*, data to extract (example : summary, customfield, description, etc...)
                - *Value 2*, type of data (example : multiplevalue, link, etc...)

                    *example : {0: ['', 'summary'], 1: ['', 'summary'], 2: ['', 'description']}*

        """

        jira_import = []

        if len(jira_issues) != 0:

            for incremental, all_issues in enumerate(jira_issues):  # Browse all issues in the JSON file

                jira_import.append({})

                for inc, table in information.items():

                    # Condition for the multiple values in the JSON variable
                    if table[0] == "multiplevalues":  # To check that it is called in the XML file

                        Fulltext = ""  # Variable that adds the names of the multiple values
                        for multiplevalues in eval("all_issues.fields." + table[1]):
                            if Fulltext == "":
                                Fulltext = str(multiplevalues.name)
                            else:
                                Fulltext = Fulltext + "\n" + str(multiplevalues.name)

                        jira_import[incremental][table[0]] = Fulltext

                    elif table[0] == "link":

                        Fulltext = []

                        for link in all_issues.fields.issuelinks:
                            if link.type.inward == table[1]:
                                if hasattr(link, "inwardIssue"):
                                    cut_string = str(link.inwardIssue.fields.summary).split(
                                        "_")  # To get the name of the issue

                                    Fulltext.append(str(cut_string[0]) + " Iss. " + str(cut_string[1]))

                            elif link.type.outward == table[1]:
                                if hasattr(link, "outwardIssue"):
                                    cut_string = str(link.outwardIssue.fields.summary).split(
                                        "_")  # To get the name of the issue

                                    Fulltext.append(str(cut_string[0]) + " Iss. " + str(cut_string[1]))

                        Fulltext = self.__just_highest_issues__(" Iss. ", 0, Fulltext)
                        Text = ""

                        for Fulltext in Fulltext:

                            if Text == "":
                                Text += Fulltext
                            else:
                                Text += "\n" + Fulltext

                        jira_import[incremental][table[1]] = Text

                    elif table[0] == "summary":
                        cut_string = str(all_issues.fields.summary).split("_")  # Split variable
                        jira_import[incremental][table[1]] = cut_string[2]

                    # For all other fields
                    else:
                        value = table[1]
                        if value in jira_import[incremental]:
                            value += str(inc)
                        jira_import[incremental][value] = str(eval("all_issues.fields." + table[1]))

        return jira_import

    def __str__(self) -> str:
        return "Jira_XMLDocument : " + self.root.attrib["name"]


if __name__ == "__main__":
    jira = JIRA(options={'server': "https://hematome.atlassian.net/"}, basic_auth=("tom.plelo.s@gmail.com", "tbjRhafn786heAzjIOVn313E"))
    jira_XML = Jira_XMLDocument(jira, "test.xml")
    jira_XML.to_excel(path="/Users/jean/Desktop")
    #jira_XML.to_word_template("CSAR-TEMPLATE.docx")
    print("finished")
